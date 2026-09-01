from __future__ import annotations

from pathlib import Path
import hashlib
import json
import math
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
OUTPUT = ROOT / "output" / "docx" / "semantic_pattern_bridge_portfolio_en.docx"
FIGURE_DIR = ROOT / "reports" / "figures"
WORK = ROOT / "tmp" / "docx" / "semantic_pattern_bridge_portfolio_academic"

SCHEMATIC = FIGURE_DIR / "pattern_semantic_parser_schematic_en.png"
DSL_EXAMPLE = FIGURE_DIR / "pattern_dsl_semantic_example_en.png"

# The portfolio figures use one real GarmentCodeData v2 source sample that was
# already included in the local benchmark corpus.  The geometry is read from
# the exact labels derived from the official specification and guarded by the
# specification SHA-256 so the diagram cannot silently fall back to invented
# coordinates.
GCDV2_FIGURE_SAMPLE_ID = "rand_LKC1OG530J"
GCDV2_FIGURE_SPEC = (
    ROOT
    / "data"
    / "processed"
    / "garmentcode_v2"
    / "batch_0_full"
    / GCDV2_FIGURE_SAMPLE_ID
    / f"{GCDV2_FIGURE_SAMPLE_ID}_specification.json"
)
GCDV2_FIGURE_LABELS = (
    ROOT
    / "artifacts"
    / "gcdv2_exact_pairs_v1"
    / "top"
    / GCDV2_FIGURE_SAMPLE_ID
    / "labels.json"
)
GCDV2_FIGURE_PATTERN = GCDV2_FIGURE_LABELS.with_name("pattern.png")

FONT_LATIN_NAME = "Arial"
FONT_EAST_ASIA_NAME = "Arial"
FONT_MONO_NAME = "Courier New"
FONT_REGULAR = Path(r"C:\Windows\Fonts\arial.ttf")
FONT_BOLD = Path(r"C:\Windows\Fonts\arialbd.ttf")

BLACK = "000000"
MUTED = "555555"
MID = "777777"
LIGHT = "EAEAEA"
LIGHTER = "F5F5F5"
ACCENT = "366091"
BLUE = "2B78B8"
ORANGE = "E47E23"
GREEN = "3A923A"
PURPLE = "744DA9"
RED = "C74747"
WHITE = "FFFFFF"

# 8.5 in Letter width - 0.55 in margins on both sides.
PAGE_WIDTH_DXA = 10656
CELL_TOP_BOTTOM_DXA = 38
CELL_START_END_DXA = 70


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_rpr_font_stack(rpr, *, mono: bool = False) -> None:
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
        rfonts.attrib.pop(qn(f"w:{key}"), None)
    latin = FONT_MONO_NAME if mono else FONT_LATIN_NAME
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), FONT_EAST_ASIA_NAME)
    rfonts.set(qn("w:cs"), latin)
    rfonts.set(qn("w:hint"), "eastAsia")

    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:eastAsia"), "ko-KR")


def set_run_font(
    run,
    *,
    size: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color: str | None = None,
    mono: bool = False,
) -> None:
    run.font.name = FONT_MONO_NAME if mono else FONT_LATIN_NAME
    set_rpr_font_stack(run._element.get_or_add_rPr(), mono=mono)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_style_font(style, *, mono: bool = False) -> None:
    style.font.name = FONT_MONO_NAME if mono else FONT_LATIN_NAME
    set_rpr_font_stack(style.element.get_or_add_rPr(), mono=mono)


def set_exact_leading(paragraph_format, points: float) -> None:
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph_format.line_spacing = Pt(points)


def configure_document_defaults(doc: Document) -> None:
    styles_element = doc.styles.element
    doc_defaults = styles_element.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles_element.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(qn("w:rPrDefault"))
    if rpr_default is None:
        rpr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(rpr_default)
    rpr = rpr_default.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        rpr_default.append(rpr)
    set_rpr_font_stack(rpr)

    settings = doc.settings.element
    theme_lang = settings.find(qn("w:themeFontLang"))
    if theme_lang is None:
        theme_lang = OxmlElement("w:themeFontLang")
        settings.append(theme_lang)
    theme_lang.set(qn("w:val"), "en-US")
    theme_lang.set(qn("w:eastAsia"), "ko-KR")


def ensure_style(doc: Document, name: str, kind=WD_STYLE_TYPE.PARAGRAPH):
    if name in doc.styles:
        return doc.styles[name]
    return doc.styles.add_style(name, kind)


def configure_styles(doc: Document) -> None:
    configure_document_defaults(doc)
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal)
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = rgb(BLACK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    set_exact_leading(normal.paragraph_format, 12.0)
    normal.paragraph_format.widow_control = True

    title = styles["Title"]
    set_style_font(title)
    title.font.size = Pt(17)
    title.font.bold = True
    title.font.color.rgb = rgb(BLACK)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    set_exact_leading(title.paragraph_format, 20)
    title.paragraph_format.keep_with_next = True
    # The built-in Word Title style may carry a theme border/shading.  This
    # report uses the reference document's plain academic title treatment.
    title_ppr = title.element.get_or_add_pPr()
    for tag in ("w:pBdr", "w:shd"):
        node = title_ppr.find(qn(tag))
        if node is not None:
            title_ppr.remove(node)

    subtitle = styles["Subtitle"]
    set_style_font(subtitle)
    subtitle.font.size = Pt(12)
    subtitle.font.italic = True
    subtitle.font.color.rgb = rgb(BLACK)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(8)
    set_exact_leading(subtitle.paragraph_format, 15)
    subtitle.paragraph_format.keep_with_next = True
    subtitle_ppr = subtitle.element.get_or_add_pPr()
    for tag in ("w:pBdr", "w:shd"):
        node = subtitle_ppr.find(qn(tag))
        if node is not None:
            subtitle_ppr.remove(node)

    h1 = styles["Heading 1"]
    set_style_font(h1)
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = rgb(BLACK)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)
    set_exact_leading(h1.paragraph_format, 18)
    h1.paragraph_format.keep_with_next = True

    h2 = styles["Heading 2"]
    set_style_font(h2)
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = rgb(BLACK)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    set_exact_leading(h2.paragraph_format, 15)
    h2.paragraph_format.keep_with_next = True

    h3 = styles["Heading 3"]
    set_style_font(h3)
    h3.font.size = Pt(9.5)
    h3.font.bold = True
    h3.font.color.rgb = rgb(BLACK)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)
    set_exact_leading(h3.paragraph_format, 12)
    h3.paragraph_format.keep_with_next = True

    portfolio = ensure_style(doc, "Portfolio Body")
    set_style_font(portfolio)
    portfolio.font.size = Pt(9.5)
    portfolio.font.color.rgb = rgb(BLACK)
    portfolio.paragraph_format.space_before = Pt(0)
    portfolio.paragraph_format.space_after = Pt(12)
    set_exact_leading(portfolio.paragraph_format, 12.7)
    portfolio.paragraph_format.widow_control = True

    meta = ensure_style(doc, "Report Meta")
    set_style_font(meta)
    meta.font.size = Pt(9)
    meta.font.color.rgb = rgb(BLACK)
    meta.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(4)
    set_exact_leading(meta.paragraph_format, 11)

    caption = ensure_style(doc, "Figure Caption")
    set_style_font(caption)
    caption.font.size = Pt(8.5)
    caption.font.bold = True
    caption.font.italic = True
    caption.font.color.rgb = rgb(BLACK)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(6)
    set_exact_leading(caption.paragraph_format, 10.5)
    caption.paragraph_format.keep_together = True

    table_caption = ensure_style(doc, "Table Caption")
    set_style_font(table_caption)
    table_caption.font.size = Pt(9.5)
    table_caption.font.bold = True
    table_caption.font.color.rgb = rgb(BLACK)
    table_caption.paragraph_format.space_before = Pt(8)
    table_caption.paragraph_format.space_after = Pt(2)
    set_exact_leading(table_caption.paragraph_format, 12)
    table_caption.paragraph_format.keep_with_next = True

    table_note = ensure_style(doc, "Table Note")
    set_style_font(table_note)
    table_note.font.size = Pt(8)
    table_note.font.italic = True
    table_note.font.color.rgb = rgb(BLACK)
    table_note.paragraph_format.space_before = Pt(0)
    table_note.paragraph_format.space_after = Pt(3)
    set_exact_leading(table_note.paragraph_format, 10)
    table_note.paragraph_format.keep_with_next = True

    table_text = ensure_style(doc, "Table Text")
    set_style_font(table_text)
    table_text.font.size = Pt(7.5)
    table_text.font.color.rgb = rgb(BLACK)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    set_exact_leading(table_text.paragraph_format, 9.3)

    table_header = ensure_style(doc, "Table Header")
    set_style_font(table_header)
    table_header.font.size = Pt(7.5)
    table_header.font.bold = True
    table_header.font.color.rgb = rgb(BLACK)
    table_header.paragraph_format.space_before = Pt(0)
    table_header.paragraph_format.space_after = Pt(0)
    set_exact_leading(table_header.paragraph_format, 9.3)

    code = ensure_style(doc, "Code Block")
    set_style_font(code, mono=True)
    code.font.size = Pt(8)
    code.font.color.rgb = rgb(BLACK)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    set_exact_leading(code.paragraph_format, 9.5)

    reference = ensure_style(doc, "Reference")
    set_style_font(reference)
    reference.font.size = Pt(9)
    reference.font.color.rgb = rgb(BLACK)
    reference.paragraph_format.left_indent = Inches(0.25)
    reference.paragraph_format.first_line_indent = Inches(-0.25)
    reference.paragraph_format.space_before = Pt(0)
    reference.paragraph_format.space_after = Pt(4)
    set_exact_leading(reference.paragraph_format, 11.2)

    appendix = ensure_style(doc, "Appendix Title")
    set_style_font(appendix)
    appendix.font.size = Pt(17)
    appendix.font.bold = True
    appendix.font.color.rgb = rgb(ACCENT)
    appendix.paragraph_format.space_before = Pt(0)
    appendix.paragraph_format.space_after = Pt(12)
    set_exact_leading(appendix.paragraph_format, 20)
    appendix.paragraph_format.keep_with_next = True

    bullet = styles["List Bullet"]
    set_style_font(bullet)
    bullet.font.size = Pt(9.5)
    bullet.font.color.rgb = rgb(BLACK)
    bullet.paragraph_format.left_indent = Inches(0.25)
    bullet.paragraph_format.first_line_indent = Inches(-0.17)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(4)
    set_exact_leading(bullet.paragraph_format, 12)

    numbered = styles["List Number"]
    set_style_font(numbered)
    numbered.font.size = Pt(9.5)
    numbered.font.color.rgb = rgb(BLACK)
    numbered.paragraph_format.left_indent = Inches(0.25)
    numbered.paragraph_format.first_line_indent = Inches(-0.18)
    numbered.paragraph_format.space_before = Pt(0)
    numbered.paragraph_format.space_after = Pt(7)
    set_exact_leading(numbered.paragraph_format, 12)


def set_section_geometry(section) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.65)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.32)
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_separate, text, fld_end])
    set_run_font(run, size=9.5, color=BLACK)


def configure_footer(section) -> None:
    section.header.paragraphs[0].text = ""
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_exact_leading(paragraph.paragraph_format, 10)
    add_page_field(paragraph)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    values = {
        "top": CELL_TOP_BOTTOM_DXA,
        "start": CELL_START_END_DXA,
        "bottom": CELL_TOP_BOTTOM_DXA,
        "end": CELL_START_END_DXA,
    }
    for tag, value in values.items():
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, *, color: str = BLACK, width: int = 4) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(width))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def proportions_to_widths(proportions: list[float]) -> list[int]:
    total = sum(proportions)
    widths = [round(PAGE_WIDTH_DXA * value / total) for value in proportions]
    widths[-1] += PAGE_WIDTH_DXA - sum(widths)
    return widths


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(PAGE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Twips(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def fill_cell(cell, text: str, *, header: bool = False, center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "Table Header" if header else "Table Text"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(text)


def add_table(
    doc: Document,
    rows: list[list[str]],
    proportions: list[float],
    *,
    center_columns: set[int] | None = None,
) -> None:
    center_columns = center_columns or set()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = proportions_to_widths(proportions)
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_header(table.rows[0])
    for r_index, row in enumerate(rows):
        for c_index, value in enumerate(row):
            cell = table.cell(r_index, c_index)
            if r_index == 0:
                set_cell_shading(cell, LIGHT)
            fill_cell(
                cell,
                value,
                header=r_index == 0,
                center=r_index == 0 or c_index in center_columns,
            )


def add_table_title(doc: Document, number: int, title: str, note: str | None = None) -> None:
    paragraph = doc.add_paragraph(style="Table Caption")
    paragraph.add_run(f"Table {number}. {title}")
    if note:
        note_paragraph = doc.add_paragraph(style="Table Note")
        note_paragraph.add_run(note)


def add_body(
    doc: Document,
    text: str,
    *,
    style: str = "Normal",
    keep_together: bool = False,
) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.keep_together = keep_together
    paragraph.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.add_run(text)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_spacer(doc: Document, points: float) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_exact_leading(paragraph.paragraph_format, points)
    paragraph.add_run("")


def add_caption(doc: Document, number: int, text: str) -> None:
    paragraph = doc.add_paragraph(style="Figure Caption")
    paragraph.add_run(f"Figure {number}. {text}")


def set_image_alt_text(shape, description: str) -> None:
    shape._inline.docPr.set("descr", description)


def add_centered_image(doc: Document, path: Path, width_inches: float, alt_text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(0)
    # Inline pictures occupy the paragraph's line box.  Inheriting Normal's
    # exact 12 pt leading clips the picture to a thin strip in LibreOffice and
    # Word.  A single, auto-height line preserves the full image.
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.keep_with_next = True
    shape = paragraph.add_run().add_picture(str(path), width=Inches(width_inches))
    set_image_alt_text(shape, alt_text)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [PAGE_WIDTH_DXA])
    set_table_borders(table, color=LIGHT, width=2)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHTER)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "Code Block"
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, size=8, mono=True, color=BLACK)


def wrapped_lines(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_multiline(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.FreeTypeFont,
    *,
    fill: str = "#000000",
    spacing: int = 8,
) -> None:
    x0, y0, x1, y1 = box
    lines = wrapped_lines(draw, text, font, x1 - x0 - 24)
    line_height = draw.textbbox((0, 0), "Ag", font=font)[3] + spacing
    height = line_height * len(lines) - spacing
    y = y0 + (y1 - y0 - height) / 2
    for line in lines:
        width = draw.textbbox((0, 0), line, font=font)[2]
        draw.text(((x0 + x1 - width) / 2, y), line, font=font, fill=fill)
        y += line_height


def cubic_points(p0, p1, p2, p3, steps: int = 40):
    points = []
    for index in range(steps + 1):
        t = index / steps
        u = 1 - t
        x = u**3 * p0[0] + 3 * u**2 * t * p1[0] + 3 * u * t**2 * p2[0] + t**3 * p3[0]
        y = u**3 * p0[1] + 3 * u**2 * t * p1[1] + 3 * u * t**2 * p2[1] + t**3 * p3[1]
        points.append((round(x), round(y)))
    return points


def load_gcdv2_figure_sample() -> dict:
    """Load and verify the exact geometry used by the portfolio figures."""

    source_digest = hashlib.sha256(GCDV2_FIGURE_SPEC.read_bytes()).hexdigest()
    labels = json.loads(GCDV2_FIGURE_LABELS.read_text(encoding="utf-8"))
    if labels.get("sample_id") != GCDV2_FIGURE_SAMPLE_ID:
        raise RuntimeError("Unexpected GCDv2 figure sample id")
    if labels.get("source_specification_sha256", "").lower() != source_digest.lower():
        raise RuntimeError("GCDv2 figure labels do not match the source specification")
    return labels


def parse_control_xy(value) -> tuple[float, float]:
    if isinstance(value, str):
        x, y = value.split()
        return float(x), float(y)
    return float(value[0]), float(value[1])


def sample_source_edge(edge: dict, *, steps: int = 48) -> list[tuple[float, float]]:
    """Sample one exact GCDv2 line or Bezier edge in source centimetres."""

    p0 = tuple(float(v) for v in edge["start_cm"])
    p3 = tuple(float(v) for v in edge["end_cm"])
    kind = edge["curve"]["type"]
    controls = [parse_control_xy(value) for value in edge["curve"].get("controls_cm", [])]
    if kind == "line":
        return [p0, p3]
    if kind == "quadratic_bezier" and len(controls) == 1:
        p1 = controls[0]
        points = []
        for index in range(steps + 1):
            t = index / steps
            u = 1 - t
            points.append(
                (
                    u * u * p0[0] + 2 * u * t * p1[0] + t * t * p3[0],
                    u * u * p0[1] + 2 * u * t * p1[1] + t * t * p3[1],
                )
            )
        return points
    if kind == "cubic_bezier" and len(controls) == 2:
        p1, p2 = controls
        points = []
        for index in range(steps + 1):
            t = index / steps
            u = 1 - t
            points.append(
                (
                    u**3 * p0[0] + 3 * u**2 * t * p1[0] + 3 * u * t**2 * p2[0] + t**3 * p3[0],
                    u**3 * p0[1] + 3 * u**2 * t * p1[1] + 3 * u * t**2 * p2[1] + t**3 * p3[1],
                )
            )
        return points
    if kind == "circular_arc":
        arc = edge["curve"].get("arc", {})
        radius = float(arc.get("radius_cm", edge["curve"]["source_params"][0]))
        chord_x = p3[0] - p0[0]
        chord_y = p3[1] - p0[1]
        chord = math.hypot(chord_x, chord_y)
        if chord <= 1e-9 or radius < chord / 2:
            raise RuntimeError("Invalid circular arc in GCDv2 figure source")

        midpoint = ((p0[0] + p3[0]) / 2, (p0[1] + p3[1]) / 2)
        height = math.sqrt(max(radius * radius - (chord / 2) ** 2, 0.0))
        normal = (-chord_y / chord, chord_x / chord)
        centers = [
            (midpoint[0] + normal[0] * height, midpoint[1] + normal[1] * height),
            (midpoint[0] - normal[0] * height, midpoint[1] - normal[1] * height),
        ]
        expected_tangent = math.radians(float(edge["start_tangent_deg"]))
        expected_large = bool(arc.get("large_arc", False))
        candidates = []
        for center in centers:
            start_angle = math.atan2(p0[1] - center[1], p0[0] - center[0])
            end_angle = math.atan2(p3[1] - center[1], p3[0] - center[0])
            for direction in (1, -1):
                if direction > 0:
                    delta = (end_angle - start_angle) % (2 * math.pi)
                else:
                    delta = -((start_angle - end_angle) % (2 * math.pi))
                tangent = start_angle + direction * math.pi / 2
                tangent_error = abs(
                    math.atan2(
                        math.sin(tangent - expected_tangent),
                        math.cos(tangent - expected_tangent),
                    )
                )
                large_error = 0.0 if (abs(delta) > math.pi) == expected_large else math.pi
                candidates.append((tangent_error + large_error, center, start_angle, delta))
        _, center, start_angle, delta = min(candidates, key=lambda item: item[0])
        return [
            (
                center[0] + radius * math.cos(start_angle + delta * index / steps),
                center[1] + radius * math.sin(start_angle + delta * index / steps),
            )
            for index in range(steps + 1)
        ]
    raise RuntimeError(f"Unsupported source curve in figure: {kind}")


def panel_transform(panel: dict, box: tuple[int, int, int, int]):
    points = [point for edge in panel["edges"] for point in sample_source_edge(edge)]
    min_x = min(point[0] for point in points)
    max_x = max(point[0] for point in points)
    min_y = min(point[1] for point in points)
    max_y = max(point[1] for point in points)
    x0, y0, x1, y1 = box
    scale = min((x1 - x0) / max(max_x - min_x, 1e-9), (y1 - y0) / max(max_y - min_y, 1e-9))
    offset_x = (x0 + x1 - (max_x - min_x) * scale) / 2
    offset_y = (y0 + y1 - (max_y - min_y) * scale) / 2

    def transform(point: tuple[float, float]) -> tuple[int, int]:
        return (
            round(offset_x + (point[0] - min_x) * scale),
            round(offset_y + (max_y - point[1]) * scale),
        )

    return transform


def draw_exact_gcdv2_panel(
    draw: ImageDraw.ImageDraw,
    panel: dict,
    box: tuple[int, int, int, int],
    *,
    semantic_roles: dict[int, str] | None = None,
) -> dict[int, tuple[int, int]]:
    """Draw one exact source panel; only screen placement and scale change."""

    transform = panel_transform(panel, box)
    edge_points: list[list[tuple[int, int]]] = []
    outline: list[tuple[int, int]] = []
    for edge in panel["edges"]:
        sampled = [transform(point) for point in sample_source_edge(edge)]
        edge_points.append(sampled)
        outline.extend(sampled if not outline else sampled[1:])
    draw.polygon(outline, fill="#E8E8E8")

    role_colors = {
        "center_front": "#D07A1F",
        "neckline": "#D33F72",
        "shoulder": "#238A6A",
        "armhole": "#744DA9",
        "side_seam": "#2B78B8",
        "hem": "#5F5F5F",
    }
    for index, sampled in enumerate(edge_points):
        role = semantic_roles.get(index) if semantic_roles else None
        color = role_colors.get(role, "#111111")
        draw.line(sampled, fill=color, width=8 if role else 5, joint="curve")

    vertices = {
        index: transform((float(point[0]), float(point[1])))
        for index, point in enumerate(panel["vertices_cm"])
    }
    for x, y in vertices.values():
        draw.ellipse((x - 6, y - 6, x + 6, y + 6), fill="#FFFFFF", outline="#111111", width=3)
    return vertices


def paste_exact_panel_set(image: Image.Image, box: tuple[int, int, int, int]) -> None:
    """Repack the eight exact contours with display-only panel scaling.

    The source sample has intentionally oversized gathered sleeves.  The model
    never consumes this packed raster or its global layout, so the overview
    normalizes each silhouette independently while preserving every contour's
    aspect ratio and analytic geometry.
    """

    source = load_gcdv2_figure_sample()
    panels = {panel["panel_id"]: panel for panel in source["panels"]}
    draw = ImageDraw.Draw(image)
    x0, y0, x1, y1 = box
    slot_width = (x1 - x0) / 4

    def draw_silhouette(panel: dict, panel_box: tuple[int, int, int, int]) -> None:
        transform = panel_transform(panel, panel_box)
        outline: list[tuple[int, int]] = []
        for edge in panel["edges"]:
            sampled = [transform(point) for point in sample_source_edge(edge)]
            outline.extend(sampled if not outline else sampled[1:])
        color = tuple(int(channel) for channel in panel["render_color_rgb"])
        draw.polygon(outline, fill=color)

    sleeve_ids = ["right_sleeve_b", "right_sleeve_f", "left_sleeve_b", "left_sleeve_f"]
    body_ids = ["right_btorso", "right_ftorso", "left_btorso", "left_ftorso"]
    for index, panel_id in enumerate(sleeve_ids):
        center_x = x0 + (index + 0.5) * slot_width
        draw_silhouette(
            panels[panel_id],
            (round(center_x - 21), y0 + 10, round(center_x + 21), y0 + 76),
        )
    for index, panel_id in enumerate(body_ids):
        center_x = x0 + (index + 0.5) * slot_width
        draw_silhouette(
            panels[panel_id],
            (round(center_x - 26), y0 + 84, round(center_x + 26), y1 - 6),
        )


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#555555") -> None:
    x0, y0 = start
    x1, y1 = end
    draw.line((x0, y0, x1 - 18, y1), fill=color, width=5)
    draw.polygon([(x1, y1), (x1 - 24, y1 - 13), (x1 - 24, y1 + 13)], fill=color)


def draw_down_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = "#555555") -> None:
    x0, y0 = start
    x1, y1 = end
    draw.line((x0, y0, x1, y1 - 14), fill=color, width=4)
    draw.polygon([(x1, y1), (x1 - 10, y1 - 18), (x1 + 10, y1 - 18)], fill=color)


def make_core_schematic(path: Path) -> None:
    source = load_gcdv2_figure_sample()
    width, height = 1800, 800
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    heading = ImageFont.truetype(str(FONT_BOLD), 30)
    body = ImageFont.truetype(str(FONT_REGULAR), 23)
    small = ImageFont.truetype(str(FONT_REGULAR), 20)
    metric = ImageFont.truetype(str(FONT_BOLD), 22)

    stages = [(30, 90, 330, 700), (400, 90, 790, 700), (860, 90, 1300, 700), (1370, 90, 1770, 700)]
    for box in stages:
        draw.rectangle(box, fill="#FFFFFF", outline="#444444", width=3)

    # Stage 1: the contours are the verified GCDv2 source geometry for the
    # selected sample.  Only non-overlap packing, scale and colour are changed.
    draw.text((180, 118), "Garment Panel Set", font=heading, fill="#111111", anchor="ma")
    draw.text((180, 160), "verified GCDv2 geometry", font=small, fill="#444444", anchor="ma")
    draw.rectangle((52, 190, 308, 475), fill="#0B0D10", outline="#777777", width=2)
    paste_exact_panel_set(image, (61, 199, 299, 466))
    panel_count = len(source["panels"])
    edge_count = sum(len(panel["edges"]) for panel in source["panels"])
    draw.text((180, 512), f"{panel_count} panels · {edge_count} edges", font=metric, fill="#111111", anchor="ma")
    draw.text((180, 550), GCDV2_FIGURE_SAMPLE_ID, font=small, fill="#444444", anchor="ma")
    draw.text((180, 592), "one analytic boundary", font=body, fill="#333333", anchor="ma")
    draw.text((180, 624), "per source panel", font=body, fill="#333333", anchor="ma")
    draw.text((180, 662), "display scale normalized", font=small, fill="#555555", anchor="ma")

    # Stage 2: only panel-local analytic geometry and cyclic relations enter
    # the neural model. Semantic labels and stitches are deliberately absent.
    draw.text((595, 118), "Analytic Edge Tokens", font=heading, fill="#111111", anchor="ma")
    draw.text((595, 166), "one ordered cycle per panel", font=small, fill="#444444", anchor="ma")
    tokens = [
        "L / Q / C / A command",
        "length · tangent · turn",
        "curve control / arc params",
        "cyclic prev · self · next",
        "no global x/y or layout",
    ]
    for index, token in enumerate(tokens):
        y = 220 + index * 82
        draw.rectangle((425, y, 765, y + 50), fill="#F0F0F0", outline="#999999", width=2)
        draw.text((595, y + 25), token, font=body, fill="#111111", anchor="mm")

    # Stage 3: replace the previous generic node rows with the actual learned
    # hierarchy. It parses an existing pattern; it does not generate geometry.
    draw.text((1080, 118), "Learned Semantic Parser", font=heading, fill="#111111", anchor="ma")
    draw.text((1080, 162), "950,820-parameter Transformer", font=small, fill="#444444", anchor="ma")
    parser_blocks = [
        (220, "2-layer edge encoder", "within each panel cycle"),
        (340, "masked mean pooling", "edge tokens → panel token"),
        (460, "3-layer garment-set encoder", "compares all panel tokens"),
        (580, "prediction heads", "category · panel · edge · seam"),
    ]
    for index, (y, label, detail) in enumerate(parser_blocks):
        draw.rectangle((890, y, 1270, y + 78), fill="#F5F7FA", outline="#7D8CA0", width=2)
        draw.text((1080, y + 22), label, font=metric, fill="#111111", anchor="ma")
        draw.text((1080, y + 53), detail, font=small, fill="#444444", anchor="ma")
        if index < len(parser_blocks) - 1:
            draw_down_arrow(draw, (1080, y + 82), (1080, y + 112), color="#6B7785")

    # Stage 4: neural scores are converted into a constrained semantic result.
    # Landmarks are junction facts at existing vertices, not coordinate heads.
    draw.text((1570, 118), "Semantic Result", font=heading, fill="#111111", anchor="ma")
    draw.text((1410, 190), "Neural proposals", font=metric, fill="#111111")
    draw.text((1410, 228), "category · panel role · edge role", font=small, fill="#444444")
    draw.text((1410, 264), "command class · seam-pair scores", font=small, fill="#444444")
    draw.line((1410, 315, 1730, 315), fill="#B0B0B0", width=2)
    draw.text((1410, 350), "Symbolic projection", font=metric, fill="#111111")
    draw.text((1410, 390), "role-cycle grammar", font=body, fill="#333333")
    draw.text((1410, 432), "seam graph matching", font=body, fill="#333333")
    draw.text((1410, 474), "shared-vertex junction rules", font=body, fill="#333333")
    draw.rectangle((1410, 535, 1730, 615), fill="#F5F7FA", outline="#7D8CA0", width=2)
    draw.text((1570, 558), "FNP / BNP / SNP / SP", font=metric, fill="#111111", anchor="ma")
    draw.text((1570, 590), "derived at existing vertices", font=small, fill="#444444", anchor="ma")
    draw.text((1570, 652), "parses patterns; does not generate them", font=small, fill="#333333", anchor="ma")

    draw_arrow(draw, (340, 395), (388, 395))
    draw_arrow(draw, (800, 395), (848, 395))
    draw_arrow(draw, (1310, 395), (1358, 395))

    credit = (
        f"Panel contours adapted from GarmentCodeData v2, sample {GCDV2_FIGURE_SAMPLE_ID} "
        "(Korosteleva et al., CC BY 4.0). Geometry preserved; placement, per-panel display scale, colour and labels changed."
    )
    draw.text((900, 752), credit, font=small, fill="#555555", anchor="mm")

    image.save(path, format="PNG", optimize=True)


def make_dsl_example(path: Path) -> None:
    source = load_gcdv2_figure_sample()
    panel = next(panel for panel in source["panels"] if panel["panel_id"] == "left_ftorso")
    width, height = 1800, 670
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    heading = ImageFont.truetype(str(FONT_BOLD), 31)
    body = ImageFont.truetype(str(FONT_REGULAR), 24)
    small = ImageFont.truetype(str(FONT_BOLD), 22)

    draw.text((275, 42), "A. Exact Source Panel", font=heading, fill="#111111", anchor="ma")
    draw_exact_gcdv2_panel(draw, panel, (105, 98, 445, 550))
    draw.text((275, 580), f"{GCDV2_FIGURE_SAMPLE_ID} · left_ftorso", font=body, fill="#333333", anchor="ma")
    draw.text((275, 614), "canonical local display frame", font=body, fill="#555555", anchor="ma")

    draw.text((890, 42), "B. Exact Analytic Cycle", font=heading, fill="#111111", anchor="ma")
    sequence = [
        "v0 ─ LINE(hem) ─ v1",
        "v1 ─ LINE(side seam) ─ v2",
        "v2 ─ CUBIC(armhole 1) ─ v3",
        "v3 ─ CUBIC(armhole 2) ─ v4",
        "v4 ─ LINE(shoulder) ─ v5",
        "v5 ─ ARC(neckline) ─ v6",
        "v6 ─ LINE(center front) ─ v0",
    ]
    for index, token in enumerate(sequence):
        y = 102 + index * 64
        draw.rectangle((625, y, 1155, y + 44), fill="#F1F1F1", outline="#999999", width=2)
        draw.text((890, y + 22), token, font=body, fill="#111111", anchor="mm")

    draw.text((1505, 42), "C. Weak Semantic Labels", font=heading, fill="#111111", anchor="ma")
    semantic_roles = {
        0: "hem",
        1: "side_seam",
        2: "armhole",
        3: "armhole",
        4: "shoulder",
        5: "neckline",
        6: "center_front",
    }
    output_vertices = draw_exact_gcdv2_panel(
        draw,
        panel,
        (1335, 98, 1675, 550),
        semantic_roles=semantic_roles,
    )
    labels = [
        ("FNP", output_vertices[6], (-66, -5)),
        ("SNP", output_vertices[5], (-60, -18)),
        ("SP", output_vertices[4], (16, -10)),
    ]
    for label, (x, y), (dx, dy) in labels:
        draw.text((x + dx, y + dy), label, font=small, fill="#111111")
    legend = [
        ("neckline", "#D33F72"),
        ("shoulder", "#238A6A"),
        ("armhole", "#744DA9"),
        ("side seam", "#2B78B8"),
        ("center front", "#D07A1F"),
        ("hem", "#5F5F5F"),
    ]
    for index, (label, color) in enumerate(legend):
        y = 572 + (index % 2) * 31
        x = 1230 + (index // 2) * 190
        draw.line((x, y, x + 34, y), fill=color, width=7)
        draw.text((x + 43, y), label, font=ImageFont.truetype(str(FONT_REGULAR), 19), fill="#333333", anchor="lm")

    draw_arrow(draw, (525, 340), (610, 340))
    draw_arrow(draw, (1165, 340), (1250, 340))
    credit_font = ImageFont.truetype(str(FONT_REGULAR), 17)
    draw.text(
        (900, 648),
        f"Exact contour: GCDv2 v2 sample {GCDV2_FIGURE_SAMPLE_ID} (CC BY 4.0). Geometry preserved; display and weak semantic labels added.",
        font=credit_font,
        fill="#555555",
        anchor="mm",
    )
    image.save(path, format="PNG", optimize=True)


def make_parser_chart(path: Path) -> None:
    width, height = 1600, 720
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    title = ImageFont.truetype(str(FONT_REGULAR), 34)
    tick = ImageFont.truetype(str(FONT_REGULAR), 24)
    value_font = ImageFont.truetype(str(FONT_BOLD), 26)
    label = ImageFont.truetype(str(FONT_REGULAR), 26)
    values = [
        ("Panel role", 0.930, BLUE),
        ("Edge role", 0.942, BLUE),
        ("Landmark", 0.928, PURPLE),
        ("Seam", 0.593, ORANGE),
    ]
    left, top, right, bottom = 165, 110, 1515, 600
    draw.text((width // 2, 34), "Canonical Pattern DSL parser · frozen test 198", font=title, fill="#111111", anchor="ma")
    for t in (0.0, 0.25, 0.5, 0.75, 1.0):
        y = bottom - int(t * (bottom - top))
        draw.line((left, y, right, y), fill="#D8D8D8", width=2)
        draw.text((left - 18, y), f"{t:.2f}", font=tick, fill="#333333", anchor="rm")
    bar_width = 190
    gap = 110
    x = 260
    for name, value, color in values:
        height_px = int(value * (bottom - top))
        y0 = bottom - height_px
        draw.rectangle((x, y0, x + bar_width, bottom), fill=f"#{color}")
        draw.text((x + bar_width / 2, y0 - 12), f"{value:.3f}", font=value_font, fill="#111111", anchor="ms")
        draw.text((x + bar_width / 2, bottom + 44), name, font=label, fill="#111111", anchor="ma")
        x += bar_width + gap
    image.save(path, format="PNG", optimize=True)


def make_retrieval_chart(path: Path) -> None:
    width, height = 1600, 720
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    title = ImageFont.truetype(str(FONT_REGULAR), 34)
    tick = ImageFont.truetype(str(FONT_REGULAR), 24)
    value_font = ImageFont.truetype(str(FONT_BOLD), 26)
    label = ImageFont.truetype(str(FONT_REGULAR), 26)
    left, top, right, bottom = 165, 110, 1515, 600
    draw.text((width // 2, 34), "Visual-to-DSL retrieval: target-matching primitive-cycle signature", font=title, fill="#111111", anchor="ma")
    for t in (0.0, 0.1, 0.2, 0.3, 0.4, 0.5):
        y = bottom - int((t / 0.5) * (bottom - top))
        draw.line((left, y, right, y), fill="#D8D8D8", width=2)
        draw.text((left - 18, y), f"{t * 100:.0f}%", font=tick, fill="#333333", anchor="rm")
    groups = [("Topology@1", 0.1465, 0.1566), ("Topology@10", 0.3636, 0.4646)]
    centers = [535, 1125]
    for center, (name, raw, trained) in zip(centers, groups):
        for x, value, color in ((center - 125, raw, "888888"), (center + 15, trained, BLUE)):
            height_px = int((value / 0.5) * (bottom - top))
            y0 = bottom - height_px
            draw.rectangle((x, y0, x + 110, bottom), fill=f"#{color}")
            draw.text((x + 55, y0 - 12), f"{value * 100:.1f}%", font=value_font, fill="#111111", anchor="ms")
        draw.text((center, bottom + 44), name, font=label, fill="#111111", anchor="ma")
    draw.rectangle((1100, 80, 1130, 104), fill="#888888")
    draw.text((1142, 92), "Raw FPN", font=tick, fill="#111111", anchor="lm")
    draw.rectangle((1320, 80, 1350, 104), fill=f"#{BLUE}")
    draw.text((1362, 92), "Trained DSL", font=tick, fill="#111111", anchor="lm")
    image.save(path, format="PNG", optimize=True)


def add_portfolio_note(doc: Document) -> None:
    add_spacer(doc, 13)
    title = doc.add_paragraph(style="Portfolio Body")
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run("Portfolio Note — 2D Sewing-Pattern Semantics")
    set_run_font(run, size=9.5, bold=True, color=BLACK)

    meta = doc.add_paragraph(style="Portfolio Body")
    meta.paragraph_format.space_after = Pt(12)
    meta.add_run("game-garment-benchmark | dataset preprocessing · pretrained-model evaluation · structured pattern learning")

    paragraphs = [
        "This portfolio does not claim to generate sewing patterns directly from four garment views. It reproduces inference from released pretrained models, decomposes their failures on out-of-domain game characters, and turns the first verifiable subproblem - reading the semantic structure of a completed 2D vector pattern - into an independent learning task.",
        "The work began by evaluating released ReWeaver and Garment Particles checkpoints. On a bounded set of eight complete official GCD-TS samples, ReWeaver reached mean panel IoU 0.8004, confirming the installation and default inference path. On SynBody and Blender characters, however, it repeatedly produced over-segmented panels and open boundaries. Garment Particles produced non-empty panel and seam structures, but manufacturability could not be established because the target samples lacked sewing ground truth and simulation validation.",
        "Rather than reduce these observations to a single image-similarity score, I represented panel roles, boundary semantics, shared vertices, landmarks, and seam relations in a panel-local analytic DSL independent of global canvas placement. I built a 1,983-pattern top/skirt/pants subset from one official GarmentCodeData v2 batch, then trained a compact Transformer and symbolic verifier to identify garment and panel roles, classify neckline, shoulder, armhole, and side-seam boundaries, and recover semantic junction landmarks.",
        "The report gives equal weight to successful and unsuccessful results. On the frozen test split, panel-role F1 was 0.930, edge-role F1 was 0.942, and landmark F1 was 0.928, while seam F1 remained 0.593. Visual-to-DSL training raised correct-topology coverage at rank 10 from 36.36% to 46.46%, but rank-1 accuracy remained 15.66%. The present contribution is therefore an inspectable 2D semantic intermediate representation and an empirical diagnosis of the next bottlenecks, not a finished generator.",
        "Throughout the project, I separated data provenance and licensing from technical accessibility and public redistribution rights. This document contains no source images with unresolved permission and no contact sheets; it uses only project-generated pattern diagrams and aggregate statistics. The repository retains preprocessing, manifest, evaluation, training, and verification code so the claims can be checked.",
    ]
    for text in paragraphs:
        add_body(doc, text, style="Portfolio Body")


def add_title_and_abstract(doc: Document) -> None:
    add_spacer(doc, 61)
    title = doc.add_paragraph(style="Title")
    title.add_run("Analytic 2D Sewing-Pattern Semantic Parser")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("Structural Semantics of Panels, Edges, Landmarks, and Seams")
    for text in ("game-garment-benchmark", "Portfolio Technical Report", "August 2026"):
        paragraph = doc.add_paragraph(style="Report Meta")
        paragraph.add_run(text)

    add_heading(doc, "Abstract", 1)
    add_body(
        doc,
        "Image-to-pattern systems may produce visually plausible outlines while still misclassifying panel roles, boundary order, shared vertices, or seam relations. This report therefore evaluates an intermediate task before end-to-end generation: semantic parsing of completed analytic vector patterns. From a 1,983-pattern top/skirt/pants subset of one official GarmentCodeData v2 batch, I constructed panel-local ordered L/Q/C/A programs that exclude absolute canvas coordinates, packed layout, source IDs, semantic-role labels, and stitch targets from neural input. On a 198-garment sample-ID-disjoint test split from the same generator, a 950,820-parameter Transformer obtained category accuracy 1.000, panel-role macro-F1 0.930, raw edge-role macro-F1 0.942, post-projection role-junction landmark F1 0.928, and post-projection seam F1 0.593. In a separate visual-to-DSL retrieval experiment, coverage of a target-matching primitive-cycle topology signature within the top ten candidates increased from 36.36% for a parameter-free raw-FPN nearest-neighbour baseline to 46.46% for a trained dual encoder, whereas top-1 changed only from 14.65% to 15.66%. A distinct 128-query four-view student, using precomputed spatial features plus known category, reduced normalized 2D semantic-coordinate MAE by 7.98% relative to a train-only category-mean baseline on 78 same-generator held-out garments. These are component-level, same-domain results; they do not establish expert drafting accuracy, cross-source generalization, end-to-end pattern generation, simulation readiness, or superiority to existing image-to-pattern systems.",
    )
    add_centered_image(
        doc,
        SCHEMATIC,
        7.35,
        f"Verified GarmentCodeData v2 sample {GCDV2_FIGURE_SAMPLE_ID} shown as eight non-overlapping source-geometry panels, converted to panel-local analytic edge tokens and parsed into constrained semantic roles, seam pairs, and role-junction landmarks.",
    )


def add_introduction(doc: Document) -> None:
    add_heading(doc, "1. Introduction", 1)
    add_body(
        doc,
        "The project began with a practical question: can front, back, left, and right views of a game character yield a 2D sewing pattern for a visually similar garment? To test available baselines, I adapted the authors' publicly accessible ReWeaver repository and released weights, and the MIT-licensed Garment Particles code and weights, to Windows on an RTX 5070 12 GB. I then ran local inference on official evaluation data and on SynBody, MPFB, and Blender characters [1,2].",
    )
    add_body(
        doc,
        "ReWeaver achieved mean panel IoU 0.8004 on the first eight complete samples sequentially extracted from the official GCD-TS archive. This was close to the paper's reported 0.8221, but it was a bounded smoke test of installation, weights, and the default inference path rather than a full test-set reproduction. Applying the same checkpoint to out-of-domain characters led to excessive panel counts and open boundaries. I therefore treated the failure as a structural generalization problem caused by differences in training distribution and input representation, rather than as a failure to run the model.",
    )
    add_body(
        doc,
        "Early repair experiments made this distinction clearer. Independent point moves and contour smoothing could reduce self-intersections and closure gaps, but they could not recover which panel was the front bodice, which curve was the armhole, or which edges should be sewn together. Treating geometry, topology, relations, and validity as one pixel loss obscures both the type of error and the module responsible for it.",
    )
    add_body(
        doc,
        "I narrowed the research question to two stages. First, can a model read the structure of a completed vector pattern without relying on human conventions for absolute placement? Second, can that semantic space later serve as the interface between four-view observations and a pattern editor? The first question is the main empirical contribution of this report. The second is kept separate as a limited preliminary experiment with a different checkpoint.",
    )
    add_body(doc, "The specific contributions are:")
    add_numbered(doc, "A panel-local Pattern DSL corpus that converts 1,983 GCDv2 vector patterns into L/Q/C/A commands, tangents, turns, and local connectivity propositions, independent of global x/y placement and packed layout.")
    add_numbered(doc, "A 0.95M-parameter Transformer parser and symbolic role-cycle verifier that jointly read panel roles, edge semantics, junction landmarks, and seam scores.")
    add_numbered(doc, "Separate frozen-split evaluations of the parser, visual-to-DSL retrieval, and symbolic projection, with topology and seam bottlenecks reported alongside successful components.")
    add_numbered(doc, "A reproducible claim boundary that distinguishes technical benchmarking from data redistribution rights, and same-generator performance from cross-source generalization.")


def add_related_work(doc: Document) -> None:
    add_heading(doc, "2. Related Work and Problem Formulation", 1)
    add_body(
        doc,
        "NeuralTailor, SewFormer, and PanelFormer reconstruct panel outlines and stitches from 3D point clouds or garment images [4-6]. ReWeaver reconstructs patterns from multiview appearance, while Garment Particles generates variable-topology garment structures conditioned on images [1,2]. These methods primarily create patterns from visual observations. In contrast, the present model reads an existing analytic 2D pattern and normalizes its internal semantic structure. Its role is therefore not to compete as a final decoder, but to inspect generated patterns and provide a structured interface for retrieval and editing modules.",
    )
    add_body(
        doc,
        "Human-image generation has widely used explicit structural conditions such as skeletons, MANO/SMPL, semantic part maps, depth and normal maps, and scene graphs [7-10]. This project does not directly use those models. It adopts the same design principle - represent verifiable symbolic structure explicitly instead of asking pixels alone to encode every relation - and applies it to panels, boundaries, landmarks, and seams.",
    )
    add_table_title(doc, 1, "Related problems and project scope")
    add_table(
        doc,
        [
            ["Problem", "Input", "Output", "Relation to this project"],
            ["Image-to-pattern reconstruction", "RGB / 4-view / point cloud", "New panel outlines and stitches", "Existing baseline; can produce inputs for this parser"],
            ["Vector-pattern semantic parsing", "Completed analytic panel graph", "Panel, edge, landmark, and seam semantics", "Current core implementation"],
            ["Pattern retrieval", "Four-view visual feature", "Existing DSL candidates from a train bank", "Implemented; does not generate a new pattern"],
            ["Parametric drafting edit", "Target semantics + basic block", "Topology-preserving residual edit", "Provisional BasicBlock pilot; no CAD or simulation validity"],
        ],
        [1.55, 1.75, 2.2, 2.5],
    )
    add_heading(doc, "2.1 Input and output contract", 2)
    add_body(
        doc,
        "One training sample contains all individual panels that compose a garment. Each panel is a closed boundary with an ordered vertex sequence and analytic edges between consecutive vertices. The model does not receive the source sample ID, packed-canvas position, absolute x/y coordinates, or ground-truth role strings. It therefore cannot exploit shortcuts such as a front panel always being placed on the left or a sleeve always being rasterized at a fixed size.",
    )
    add_body(
        doc,
        "Outputs are the garment category, a weak semantic role for each source panel, a semantic role for each edge, a landmark set derived from vertices shared by predicted edge roles, and scores for possible stitch partners. Labels such as FNP, SNP, and SP are not regressed as independent coordinates. They are defined as edge-role junctions, for example center front meeting neckline or neckline meeting shoulder. This preserves the real pattern structure in which multiple boundary segments share one vertex.",
    )


def add_dataset_representation(doc: Document) -> None:
    add_heading(doc, "3. Dataset and Analytic Representation", 1)
    add_body(
        doc,
        "The primary data source is the official vector sewing-pattern specification in GarmentCodeData v2 [3]. It contains panel vertices, line, quadratic/cubic Bezier and arc parameters, stitch pairs, simulated 3D drapes, and renders. The repository does not redistribute the complete dataset. It retains only lightweight manifests, splits, and canonical tokens extracted from one official batch for training and verification.",
    )
    add_table_title(doc, 2, "Canonical Pattern DSL corpus")
    add_table(
        doc,
        [
            ["Item", "Count / setting", "Interpretation"],
            ["Canonical vector patterns", "1,983", "Top/skirt/pants subset from one official batch; independent of global x/y and packed layout"],
            ["Train / validation / frozen test", "1,587 / 198 / 198", "Sample-ID disjoint; within the same GCDv2 generator"],
            ["Exact pattern-four-view intersection", "1,962", "Project-generated Blender views: top 923 / skirt 780 / pants 259"],
            ["Visual retrieval train bank", "1,570", "Existing DSL library searched by test queries"],
            ["Expert-approved cross-source test", "0", "Most important current validation gap"],
        ],
        [2.15, 2.05, 3.8],
        center_columns={1},
    )
    add_heading(doc, "3.1 Panel-local canonical program", 2)
    add_body(
        doc,
        "For each panel, the first vertex and traversal direction are canonicalized, and every boundary edge is stored as an L/Q/C/A primitive. Lengths and Bezier control points are normalized in a panel-local chord frame. Start and end tangents, turns into the next edge, and NEXT and SHARED_ENDPOINT facts form the neural input. SEWN_TO is retained only as supervision for the seam head. The representation emphasizes which segment follows which curve and which vertices are shared, rather than where the panel happens to sit on a page.",
    )
    add_centered_image(
        doc,
        DSL_EXAMPLE,
        7.0,
        f"Exact left-front bodice panel from GarmentCodeData v2 sample {GCDV2_FIGURE_SAMPLE_ID} represented as a seven-edge analytic cycle and annotated with weak semantic edge roles and FNP/SNP/SP junction landmarks.",
    )
    add_caption(
        doc,
        1,
        f"Exact left-front bodice panel from GCDv2 sample {GCDV2_FIGURE_SAMPLE_ID}. The center lists its seven source line, cubic-Bezier, and circular-arc commands; colors and FNP/SNP/SP on the right are automatically derived weak semantic labels. Geometry is unchanged; only layout, scale, color, and labels are adapted.",
    )
    add_heading(doc, "3.2 Label provenance", 2)
    add_body(
        doc,
        "Panel and edge semantics were derived automatically from GCDv2 topology, panel names, stitch relations, boundary order, and a rule resolver. These labels are consistent weak ground truth within the same generator; they are not an industry-standard ontology approved by a pattern-making expert. Seam allowance, notches, grainlines, and grading rules were excluded because the source does not provide definitive targets for them.",
    )


def add_model_and_protocol(doc: Document) -> None:
    add_heading(doc, "4. Model and Evaluation Protocol", 1)
    add_body(
        doc,
        "The canonical parser is a 0.95M-parameter Transformer that processes both cyclic primitive tokens within each panel and the full set of panels in a garment. Each primitive token contains a type embedding, normalized length, chord-frame curve parameters, start and end tangents, and turn. A panel encoder reads the ordered boundary, while a garment-set encoder compares role and stitch candidates across panels.",
    )
    add_table_title(doc, 3, "Prediction heads and supervision")
    add_table(
        doc,
        [
            ["Head", "Prediction unit", "Target / metric"],
            ["Garment category", "garment set", "cross-entropy / accuracy"],
            ["Panel role", "panel token", "role class / macro-F1"],
            ["Edge role", "boundary primitive", "neckline, armhole, side seam, etc. / macro-F1"],
            ["Landmark", "predicted role junction", "junction set / set F1"],
            ["Seam", "edge-pair score", "source stitch pair / symbolic F1"],
        ],
        [1.6, 2.0, 4.4],
    )
    add_heading(doc, "4.1 Symbolic constraints", 2)
    add_body(
        doc,
        "Edges are not treated only as independent classes. A basic front bodice, for example, has an allowed role cycle: center front -> neckline -> shoulder -> armhole -> side seam -> hem. After inference, a finite-state projection repairs impossible role orders, and a landmark is emitted only when the posteriors of two edge roles meet at the same vertex. This stage checks semantic consistency without changing geometry.",
    )
    add_heading(doc, "4.2 Frozen evaluation", 2)
    add_body(
        doc,
        "The 198 test garments were not used for model selection or early stopping. Panel, edge, and landmark performance is reported with macro-F1; seam performance uses symbolic F1 over the recovery of exact pairs among possible edge pairs. These results measure sample-ID-unseen performance within the same GarmentCode generator and renderer. Generalization to recipe-family-disjoint, body-disjoint, cross-generator, and scanned real-world patterns has not yet been measured.",
    )
    add_heading(doc, "4.3 Visual-to-DSL retrieval", 2)
    add_body(
        doc,
        "A separate 0.89M visual bridge aligns front, back, left, and right FPN features with frozen DSL embeddings. Each test query retrieves the nearest existing patterns from a bank of 1,570 training garments. The topology signature contains only the category and panel-wise cyclic L/Q/C/A sequences; it does not include exact coordinates, semantic edge roles, or the complete seam graph. Raw FPN nearest-neighbour retrieval is parameter-free whereas the DSL lane is trained, so this is not an architecture-matched ablation.",
    )


def add_results(doc: Document, parser_chart: Path, retrieval_chart: Path) -> None:
    add_heading(doc, "5. Results", 1)
    add_heading(doc, "5.1 Semantic parsing of completed vector patterns", 2)
    add_centered_image(
        doc,
        parser_chart,
        6.8,
        "Bar chart of panel role, edge role, landmark, and seam F1 on the frozen 198-garment test set.",
    )
    add_caption(
        doc,
        2,
        "Semantic parsing on 198 frozen-test garments. Panel, edge, and landmark scores were approximately 0.93, while post-projection symbolic seam F1 was materially lower at 0.593.",
    )
    add_table_title(doc, 4, "Canonical parser frozen-test metrics", "1,983 patterns; train 1,587 / validation 198 / frozen test 198.")
    add_table(
        doc,
        [
            ["Output", "Metric", "Result", "Interpretation"],
            ["Garment category", "Accuracy", "1.000", "Category classification within GCDv2"],
            ["Panel role", "Macro-F1", "0.930", "Front/back bodices, sleeves, trouser panels, and related roles"],
            ["Raw edge role", "Macro-F1", "0.942", "Neckline, shoulder, armhole, side seam, and related roles"],
            ["Projected landmark set", "F1", "0.928", "Derived from predicted edge-role junctions"],
            ["Projected symbolic seam", "F1", "0.593", "Global pair matching remains unresolved"],
        ],
        [1.65, 1.25, 1.1, 4.0],
        center_columns={1, 2},
    )
    add_body(
        doc,
        "The panel and edge metrics are not merely classifications of primitive shape. They measure panel roles and edge functions jointly within a garment set after removing absolute canvas coordinates and packed-panel location. Landmarks are not detected as independent heatmap points; they are derived from vertices where predicted roles meet. The result is consistent with the model using ordered primitives and their relations and rules out the most direct packed-canvas shortcut, although generator-specific shortcuts may remain.",
    )
    add_body(
        doc,
        "Seam recovery requires garment-level one-to-one or one-to-many matching together with length, direction, and ease constraints; local edge similarity is not sufficient. The current pair scorer does not model this global assignment well enough. The parser is therefore reported as a usable 2D semantic teacher with a clear seam bottleneck, not as a system that understands patterns like a patternmaker.",
    )

    add_heading(doc, "5.2 Visual retrieval and topology constraints", 2)
    add_centered_image(
        doc,
        retrieval_chart,
        6.8,
        "Comparison between raw FPN nearest-neighbour retrieval and a trained visual-to-DSL retriever for the target-matching primitive-cycle signature at ranks 1 and 10.",
    )
    add_caption(
        doc,
        3,
        "Raw FPN and trained visual-to-DSL retrieval. Coverage of the target-matching primitive-cycle topology signature at rank 10 rose from 36.36% to 46.46% on this fixed split, while rank-1 changed only from 14.65% to 15.66%.",
    )
    add_table_title(doc, 5, "Visual-to-DSL retrieval on the frozen 198-garment test")
    add_table(
        doc,
        [
            ["Method", "Category@1", "Topology@1", "Topology@10"],
            ["Raw FPN nearest neighbour", "99.49%", "14.65%", "36.36%"],
            ["Trained visual–DSL encoder", "98.48%", "15.66%", "46.46%"],
            ["Difference", "-1.01%p", "+1.01%p", "+10.10%p"],
        ],
        [3.3, 1.55, 1.55, 1.6],
        center_columns={1, 2, 3},
    )
    add_body(
        doc,
        "The evidence-safe interpretation is that the trained DSL retriever broadened the candidate set available to a downstream verifier or editor without reliably selecting the final target at rank 1. Category@1 was already nearly saturated for raw FPN features and fell by 1.01 percentage points after training, so the main retrieval bottleneck is distinguishing topology within a category. The comparison is not architecture-matched: the raw nearest-neighbour baseline is parameter-free, while the DSL dual encoder is trained.",
    )
    add_body(
        doc,
        "Applying symbolic projection to retrieved candidates reduced defined role-cycle violations from 168 to 0, but target-signature topology@1 remained 15.66%. This negative result separates two problems: eliminating candidates that violate the role grammar and selecting the candidate that matches the target four-view appearance. Projection did not repair geometry, prove sewability, or recover the target seam graph.",
    )


def add_fourview_pilot(doc: Document) -> None:
    add_heading(doc, "6. Preliminary Four-view Bridge", 1)
    add_body(
        doc,
        "To test whether four-view observations can enter a related semantic space, I ran a separate 128-query teacher-student experiment. It uses a different ontology and checkpoint from the 0.95M canonical parser above: a 3.23M vector-graph teacher and a 3.62M student with 9 panel, 35 path, 62 landmark, and 22 reference-line queries. The student receives four precomputed spatial feature tensors plus the known garment category and predicts query presence and normalized 2D coordinates. Seam queries are not included.",
    )
    add_body(
        doc,
        "The student does not place FNP or SP on a 3D mesh surface. It predicts where the corresponding semantic point lies in the normalized [0,1] garment-union frame of the paired 2D pattern. The score is therefore a component test of sample-specific semantic-coordinate variation beyond a category mean, not direct recovery of CAD coordinates from pixels.",
    )
    add_table_title(doc, 6, "Same-generator held-out 4-view pilot", "517 strict-common records; natural split 361 / 78 / 78.")
    add_table(
        doc,
        [
            ["Method", "Normalized coordinate MAE", "Result"],
            ["Train-only category mean", "0.046059", "Per-category mean-coordinate prior"],
            ["4-view semantic student", "0.042384", "Precomputed four-view features + known category"],
            ["Relative change", "-7.98%", "Lower error on this held-out split"],
        ],
        [3.0, 2.1, 2.9],
        center_columns={1},
    )
    add_heading(doc, "6.1 Validation-gated BasicBlock editing", 2)
    add_body(
        doc,
        "To test whether the student output could inform downstream editing, I applied residuals to a deterministic category-default provisional BasicBlock. Only queries that improved over the anchor on validation were eligible. Edits that failed topology-preserving projection or a 1.0/0.75/0.5/0.25 line search reverted to the anchor. Test ground truth was not used to choose the gate.",
    )
    add_table_title(doc, 7, "Guarded semantic editing result")
    add_table(
        doc,
        [
            ["Item", "Anchor", "After edit", "Interpretation"],
            ["Test 78 coordinate MAE", "0.087695", "0.085783", "2.18% relative reduction"],
            ["Samples improved", "—", "51 / 78", "65.4%"],
            ["Edit status", "—", "72 applied / 6 fallback", "Applied does not mean improved"],
            ["One-to-many path", "—", "99 requests rejected", "Instance identity was ambiguous"],
        ],
        [2.0, 1.45, 2.35, 2.2],
        center_columns={1, 2},
    )
    add_body(
        doc,
        "The 2.18% reduction is not an improvement in CAD geometry, sewability, fit, or simulation pass rate. It is the error after projecting the default and edited blocks into comparable semantic-query coordinates for the GCD target. The editor also uses a deterministic category-default provisional BasicBlock rather than the retrieved DSL anchor, so the parser/retrieval lane and the student/editor lane remain separate inference graphs.",
    )


def add_discussion(doc: Document) -> None:
    add_heading(doc, "7. Discussion", 1)
    add_body(
        doc,
        "The experiments show that 2D pattern semantics can function as a learnable intermediate representation rather than only as descriptive labels. Strong panel, edge, and landmark scores after removing absolute canvas placement are consistent with the use of primitive order and relations and rule out the most direct packed-layout shortcut, although generator-specific shortcuts may remain. The higher top-10 retrieval coverage and the separate four-view pilot's lower error than a category prior provide limited evidence that visual observations can be aligned with this semantic space.",
    )
    add_body(
        doc,
        "The limits are equally clear. Topology@1 of 15.66% shows that selecting the final target structure within a category remains largely unresolved. Post-projection seam F1 of 0.593 shows that strong local edge classification does not automatically recover the garment-level seam graph. Symbolic projection removes violations of the defined role-cycle grammar but cannot turn a different valid candidate into the target pattern.",
    )
    add_heading(doc, "7.1 Why the remaining errors occur", 2)
    add_bullet(doc, "A seam is a global matching problem, not an independent edge class. It must jointly satisfy lengths, directions, degree, ease, and one-to-many instances on both sides.")
    add_bullet(doc, "The primitive-cycle topology signature compresses exact coordinates, curvature, edge semantics, and inter-panel layout, so multiple patterns can share the same signature.")
    add_bullet(doc, "GCDv2-derived labels are internally consistent but are not equivalent to expert-approved industrial FNP/BNP/SNP/SP labels, seam allowances, notches, or grainlines.")
    add_bullet(doc, "The four-view pilot is sample-ID unseen but uses the same generator, neutral body, and render style. It does not establish cross-source or real-image generalization.")
    add_heading(doc, "7.2 Relation to ReWeaver and Garment Particles", 2)
    add_body(
        doc,
        "The current parser cannot be claimed as a better pattern generator than ReWeaver or Garment Particles. No same-input, same-split, same-metric comparison has been run between raw outputs and a parser-verifier-editor pipeline. A decisive next evaluation would run raw ReWeaver and ReWeaver plus the bridge on identical four-view inputs and compare closed-panel rate, seam violations, simulation pass rate, silhouette IoU, expert preference, and abstention rate.",
    )


def add_limitations_future(doc: Document) -> None:
    add_heading(doc, "8. Limitations and Future Work", 1)
    add_body(
        doc,
        "First, the frozen test is an internal split from the same GarmentCode generator. A design-family-disjoint split is needed to control more strictly for similar recipe variants appearing across train and test. FreeSewing patterns and 30-100 expert-reviewed real patterns should form a separate cross-source test.",
    )
    add_body(
        doc,
        "Second, the seam head should move from independent pair scores to bipartite graph matching with degree, direction, length, and ease as hard or soft constraints. For example, sleeve-cap length should match the sum of the front and back armholes plus specified ease. A solver should check this explicit sewing rule instead of asking the model to rediscover it unreliably from data.",
    )
    add_body(
        doc,
        "Third, target-conditioned topology selection is required. Cross-attention between four-view semantic tokens and candidate DSL tokens should separate grammatical validity from target compatibility. Once a candidate is selected, the model should predict residuals for drafting parameters such as neck width and depth, shoulder slope, armhole depth, sleeve cap, and bodice length. A parametric decoder should then redraw shared vertices jointly instead of moving raw control points independently.",
        keep_together=True,
    )
    add_table_title(doc, 8, "Next decisive experiments")
    add_table(
        doc,
        [
            ["Step", "Experiment", "Pass criterion"],
            ["1", "Target four-view token x candidate DSL cross-attention", "Material increase in topology@1"],
            ["2", "Bipartite seam matching + length/direction/ease", "Improve seam F1 and constraint pass together"],
            ["3", "Counterfactual four-view pairs with one drafting parameter changed", "Held-out regression of change direction and magnitude"],
            ["4", "Parametric drafting decoder + shared-point solver", "Closed panels, symmetry, and armhole-sleeve relation satisfied"],
            ["5", "Raw ReWeaver vs. ReWeaver + bridge A/B", "Improved simulation, silhouette, and expert metrics"],
            ["6", "FreeSewing + expert-approved real patterns", "Reported cross-source frozen-test performance"],
        ],
        [0.65, 4.55, 2.8],
        center_columns={0},
    )
    add_body(
        doc,
        "The long-term goal is not to regenerate every output from an image-to-pattern model. It is to preserve valid patterns, flag semantically impossible structures with explainable errors, and abstain from editing when confidence is low. Separating the responsibilities of generation, constraint solving, simulation, and expert review is necessary to extend this portfolio prototype into a production workflow.",
    )


def add_appendix(doc: Document) -> None:
    doc.add_page_break()
    paragraph = doc.add_paragraph(style="Appendix Title")
    paragraph.add_run("Appendix")

    add_heading(doc, "Appendix A. Reproducibility map", 2)
    add_body(
        doc,
        "The following files are the minimum evidence needed to recheck the report's core metrics and data contracts. Original datasets, checkpoints, caches, and large artifacts are excluded from Git; the public package retains code, lightweight manifests, this portfolio, and its project-generated figures.",
    )
    add_table_title(doc, 9, "Tracked evidence and implementation entry points")
    add_table(
        doc,
        [
            ["Path", "Purpose"],
            ["data/manifests/gcdv2_pattern_dsl_v2.json", "1,983-pattern canonical DSL corpus and frozen split"],
            ["data/manifests/gcdv2_pattern_dsl_v2_training_result.json", "Panel·edge·landmark·seam parser metrics"],
            ["data/manifests/gcdv2_visual_pattern_dsl_retrieval_v2.json", "Raw-FPN vs. trained-DSL retrieval"],
            ["data/manifests/basic_semantic_teacher_student_v3.json", "Separate 128-query teacher/student and guarded editor"],
            ["data/manifests/semantic_pattern_bridge_portfolio_proof.json", "Final claim boundaries and evidence index"],
            ["benchmark/scripts/train_gcdv2_pattern_dsl_transformer.py", "Canonical parser training"],
            ["benchmark/scripts/train_gcdv2_visual_dsl_retrieval.py", "Visual-to-DSL bridge training"],
            ["benchmark/scripts/train_basic_semantic_teacher_student.py", "128-query preliminary bridge"],
        ],
        [4.45, 3.55],
    )

    add_heading(doc, "Appendix B. Result boundaries", 2)
    add_table_title(doc, 10, "Result boundaries")
    add_table(
        doc,
        [
            ["Supported by the current evidence", "Not supported yet"],
            ["Learns analytic panel/edge/landmark semantics within GCDv2", "Fully recovers expert industrial-pattern semantics"],
            ["DSL alignment increases top-10 coverage of target-matching topology signatures", "Generates the exact pattern from four views or solves top-1 selection"],
            ["Separate four-view student achieves 7.98% lower normalized-coordinate MAE than the train-only category-mean baseline", "Generalizes to real photos, other generators, or bodies"],
            ["Validation-gated editing reduces query error for 51/78 samples", "Passes sewability, fit, simulation, or manufacturing checks"],
            ["Bounded ReWeaver smoke test runs on eight official samples", "Outperforms ReWeaver end to end"],
        ],
        [4.0, 4.0],
    )

    add_heading(doc, "Appendix C. Data-use boundary", 2)
    add_body(
        doc,
        "NOVA-Human is recorded with provenance_status=ARCHIVED_OFFICIAL_LINK_LIVE, license_status=UNRESOLVED, and usage_status=PRIVATE_EVALUATION_ONLY. No permission beyond private evaluation is assumed. This document contains no NOVA-Human source images, private contact sheets, live storage links, or samples with unclear redistribution rights. GarmentCodeData v2 and referenced public models require separate checks of their data, code, and checkpoint licenses. Passing a technical benchmark does not grant permission to redistribute or publicly display source data.",
    )
    add_body(
        doc,
        "GarmentCodeData v2 was created by Maria Korosteleva, Timur Levent Kesdogan, Fabian Kemper, Stephan Wenninger, Jasmin Koller, Yuhan Zhang, Mario Botsch, and Olga Sorkine-Hornung, and released under CC BY 4.0 (https://doi.org/10.3929/ethz-b-000690432; https://creativecommons.org/licenses/by/4.0/). This project parses and normalizes the official vector specification into a panel-local analytic DSL and derives semantic labels, split metadata, and aggregate statistics. The original dataset archive, renders, and meshes are excluded from this document and the public ZIP; no endorsement by the original authors is implied.",
    )


def add_references(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "References", 1)
    references = [
        "[1] ReWeaver project and released GCD-TS benchmark. Project page: https://sii-liming.github.io/ReWeaver/ ; author dataset: https://huggingface.co/datasets/SII-LiMing/ReWeaver-GCD-TS",
        "[2] Garment Particles. Official code repository: https://github.com/garment-particles/GarmentParticles ; released checkpoint repository: https://huggingface.co/georgeNakayama/GarmentParticles",
        "[3] Korosteleva, M. et al. GarmentCodeData v2: 115,000+ made-to-measure garments with sewing patterns and simulated drapes. ETH Research Collection: https://www.research-collection.ethz.ch/handle/20.500.11850/690432",
        "[4] Korosteleva, M. and Lee, S.-H. NeuralTailor: Reconstructing Sewing Pattern Structures from 3D Point Clouds of Garments. ACM Transactions on Graphics, 2022. https://github.com/maria-korosteleva/Garment-Pattern-Estimation",
        "[5] SewFormer: Sewing Pattern Reconstruction from a Single Image. Project page: https://sewformer.github.io/",
        "[6] Chen, X. et al. PanelFormer: Sewing Pattern Reconstruction from 2D Garment Images. WACV, 2024. https://openaccess.thecvf.com/content/WACV2024/html/Chen_Panelformer_Sewing_Pattern_Reconstruction_From_2D_Garment_Images_WACV_2024_paper.html",
        "[7] Narasimhaswamy, S. et al. HanDiffuser: Text-to-Image Generation with Realistic Hand Appearances. CVPR, 2024. https://openaccess.thecvf.com/content/CVPR2024/html/Narasimhaswamy_HanDiffuser_Text-to-Image_Generation_With_Realistic_Hand_Appearances_CVPR_2024_paper.html",
        "[8] Li, S. et al. CosmicMan: A Text-to-Image Foundation Model for Humans. CVPR, 2024. https://openaccess.thecvf.com/content/CVPR2024/html/Li_CosmicMan_A_Text-to-Image_Foundation_Model_for_Humans_CVPR_2024_paper.html",
        "[9] Liu, X. et al. HyperHuman: Hyper-Realistic Human Generation with Latent Structural Diffusion. ICLR, 2024. https://openreview.net/forum?id=duyA42HlCK",
        "[10] Johnson, J. et al. Image Generation from Scene Graphs. CVPR, 2018. https://openaccess.thecvf.com/content_cvpr_2018/html/Johnson_Image_Generation_From_CVPR_2018_paper.html",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(style="Reference")
        paragraph.add_run(reference)


def audit_document(doc: Document) -> None:
    if len(doc.sections) != 1:
        raise AssertionError(f"Expected one uniform section, found {len(doc.sections)}")
    section = doc.sections[0]
    if round(section.page_width.inches, 3) != 8.5 or round(section.page_height.inches, 3) != 11.0:
        raise AssertionError("Document must remain US Letter portrait")
    if round(section.left_margin.inches, 3) != 0.55 or round(section.right_margin.inches, 3) != 0.55:
        raise AssertionError("Horizontal margins do not match the reference layout")
    if round(section.top_margin.inches, 3) != 0.75:
        raise AssertionError("Top margin does not match the reference layout")
    for table in doc.tables:
        tbl_w = table._tbl.tblPr.find(qn("w:tblW"))
        if tbl_w is None or tbl_w.get(qn("w:w")) != str(PAGE_WIDTH_DXA):
            raise AssertionError("Table width is not full text width")
        widths = [int(node.get(qn("w:w"))) for node in table._tbl.tblGrid.findall(qn("w:gridCol"))]
        if sum(widths) != PAGE_WIDTH_DXA:
            raise AssertionError(f"Table grid width mismatch: {widths}")


def build() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    WORK.mkdir(parents=True, exist_ok=True)

    parser_chart = WORK / "parser_f1.png"
    retrieval_chart = WORK / "retrieval_topology.png"
    make_core_schematic(SCHEMATIC)
    make_dsl_example(DSL_EXAMPLE)
    make_parser_chart(parser_chart)
    make_retrieval_chart(retrieval_chart)

    doc = Document()
    configure_styles(doc)
    set_section_geometry(doc.sections[0])
    configure_footer(doc.sections[0])

    doc.core_properties.title = "Analytic 2D Sewing-Pattern Semantic Parser"
    doc.core_properties.subject = "Coordinate-invariant analytic pattern DSL and semantic parsing for garment inverse design"
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.comments = ""
    doc.core_properties.keywords = "sewing pattern, semantic parser, analytic DSL, garment topology, portfolio"

    add_portfolio_note(doc)
    doc.add_page_break()
    add_title_and_abstract(doc)
    doc.add_page_break()
    add_introduction(doc)
    add_related_work(doc)
    add_dataset_representation(doc)
    add_model_and_protocol(doc)
    add_results(doc, parser_chart, retrieval_chart)
    add_fourview_pilot(doc)
    add_discussion(doc)
    add_limitations_future(doc)
    add_appendix(doc)
    add_references(doc)

    audit_document(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
